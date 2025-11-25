import os
import json
from http.server import BaseHTTPRequestHandler
from io import BytesIO

from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter

from docx import Document
import csv

WORKER_KEY = os.getenv("PY_WORKER_KEY", "")


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        # ---- AUTH ----
        auth = self.headers.get("Authorization", "")
        if not auth.startswith("Bearer ") or auth.split(" ", 1)[1] != WORKER_KEY:
            self.send_response(401)
            self.end_headers()
            self.wfile.write(b"Unauthorized")
            return

        # ---- BODY ----
        length = int(self.headers.get("Content-Length", 0))
        raw = self.rfile.read(length)
        try:
            body = json.loads(raw)
        except Exception:
            body = {}

        export_type = body.get("type", "")
        title = body.get("title", "Solace Export")
        content = body.get("content", "")

        if export_type == "pdf":
            buf = self.generate_pdf(title, content)
            mime = "application/pdf"
        elif export_type == "docx":
            buf = self.generate_docx(title, content)
            mime = (
                "application/"
                "vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif export_type == "csv":
            buf = self.generate_csv(content)
            mime = "text/csv"
        else:
            self.send_response(400)
            self.end_headers()
            self.wfile.write(b"Unsupported type")
            return

        self.send_response(200)
        self.send_header("Content-Type", mime)
        self.end_headers()
        self.wfile.write(buf)

    def generate_pdf(self, title: str, text: str) -> bytes:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph(f"<b>{title}</b>", styles["Heading1"])]

        for para in text.split("\n"):
            story.append(Paragraph(para, styles["BodyText"]))

        doc.build(story)
        return buffer.getvalue()

    def generate_docx(self, title: str, text: str) -> bytes:
        doc = Document()
        doc.add_heading(title, level=1)
        for para in text.split("\n"):
            doc.add_paragraph(para)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def generate_csv(self, text: str) -> bytes:
        buffer = BytesIO()
        writer = csv.writer(buffer)
        for line in text.replace("\r", "").split("\n"):
            writer.writerow([line])
        return buffer.getvalue()
