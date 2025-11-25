import os
import io
import csv
from fastapi import FastAPI, HTTPException, Request, Response

from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document

# ------------------------------------------------------
# FastAPI app (Vercel detects this automatically)
# ------------------------------------------------------
app = FastAPI(title="MCA Export Worker")

PY_WORKER_KEY = os.getenv("PY_WORKER_KEY", "")

# ------------------------------------------------------
# Helpers
# ------------------------------------------------------
def make_pdf(title: str, text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()

    story = [Paragraph(f"<b>{title}</b>", styles["Heading1"])]

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        story.append(Paragraph(line, styles["BodyText"]))

    doc.build(story)
    return buf.getvalue()


def make_docx(title: str, text: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_csv(text: str) -> bytes:
    buf = io.StringIO()
    w = csv.writer(buf)
    for line in text.replace("\r", "").split("\n"):
        w.writerow([line])
    return buf.getvalue().encode("utf-8")


# ------------------------------------------------------
# Health endpoints
# ------------------------------------------------------
@app.get("/")
def root():
    return {"ok": True, "service": "mca-export-worker"}


@app.get("/healthz")
def healthz():
    return {"ok": True, "key_present": bool(PY_WORKER_KEY)}


# ------------------------------------------------------
# Main export endpoint
# ------------------------------------------------------
@app.post("/api/generate")
async def generate(request: Request):

    if not PY_WORKER_KEY:
        raise HTTPException(500, "PY_WORKER_KEY not set on server")

    auth = request.headers.get("authorization", "")
    if auth != f"Bearer {PY_WORKER_KEY}":
        raise HTTPException(401, "Unauthorized")

    try:
        body = await request.json()
    except:
        raise HTTPException(400, "Invalid JSON")

    export_type = (body.get("type") or body.get("format") or "").lower().strip()
    title = (body.get("title") or "Solace Export").strip()
    content = body.get("content") or ""

    if export_type not in ("pdf", "docx", "csv"):
        raise HTTPException(400, f"Invalid export type: {export_type}")

    if export_type == "pdf":
        blob = make_pdf(title, content)
        media = "application/pdf"
        fname = "solace-export.pdf"

    elif export_type == "docx":
        blob = make_docx(title, content)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        fname = "solace-export.docx"

    else:  # csv
        blob = make_csv(content)
        media = "text/csv"
        fname = "solace-export.csv"

    headers = {
        "Content-Disposition": f'inline; filename="{fname}"'
    }

    return Response(content=blob, media_type=media, headers=headers)
