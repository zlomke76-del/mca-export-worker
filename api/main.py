import os
import io
import csv

from fastapi import FastAPI, HTTPException, Request, Response
from pydantic import BaseModel

from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter

from docx import Document

app = FastAPI()

PY_WORKER_KEY = os.getenv("PY_WORKER_KEY", "")


class ExportPayload(BaseModel):
    type: str
    title: str = "Solace Export"
    content: str = ""


def generate_pdf(title: str, text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()

    story = [Paragraph(f"<b>{title}</b>", styles["Heading1"])]
    for para in text.split("\n"):
        story.append(Paragraph(para, styles["BodyText"]))

    doc.build(story)
    return buf.getvalue()


def generate_docx(title: str, text: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in text.split("\n"):
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_csv(text: str) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)

    rows = text.replace("\r", "").split("\n")
    for r in rows:
        writer.writerow([r])

    return buf.getvalue().encode("utf-8")


@app.post("/api/generate")
async def generate(request: Request, payload: ExportPayload):
    # ---- Auth check ----
    auth = request.headers.get("authorization", "")
    expected = f"Bearer {PY_WORKER_KEY}" if PY_WORKER_KEY else None

    if not expected or auth != expected:
        raise HTTPException(status_code=401, detail="Unauthorized")

    export_type = payload.type.lower().strip()

    if export_type == "pdf":
        data = generate_pdf(payload.title, payload.content)
        media_type = "application/pdf"
    elif export_type == "docx":
        data = generate_docx(payload.title, payload.content)
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif export_type == "csv":
        data = generate_csv(payload.content)
        media_type = "text/csv"
    else:
        raise HTTPException(status_code=400, detail="Unsupported type")

    return Response(content=data, media_type=media_type)
