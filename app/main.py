import os
import io
import csv
import logging
from typing import Any, Dict

from fastapi import FastAPI, HTTPException, Request, Response

from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter

from docx import Document

# ------------------------------------------------------
# Logging setup
# ------------------------------------------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("mca-export-worker")

# ------------------------------------------------------
# FastAPI app that Vercel discovers automatically
# ------------------------------------------------------
app = FastAPI(title="MCA Export Worker")

# Shared secret – must match PY_WORKER_KEY in the studio.moralclarity.ai project
PY_WORKER_KEY = os.getenv("PY_WORKER_KEY", "")


# ------------------------------------------------------
# Helpers
# ------------------------------------------------------
def generate_pdf(title: str, text: str) -> bytes:
    """Very simple PDF generator using reportlab."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()

    story = [Paragraph(f"<b>{title}</b>", styles["Heading1"])]

    for para in text.split("\n"):
        para = para.strip()
        if not para:
            continue
        story.append(Paragraph(para, styles["BodyText"]))

    doc.build(story)
    return buf.getvalue()


def generate_docx(title: str, text: str) -> bytes:
    """Very simple DOCX generator using python-docx."""
    doc = Document()
    doc.add_heading(title, level=1)

    for para in text.split("\n"):
        para = para.strip()
        if not para:
            continue
        doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_csv(text: str) -> bytes:
    """
    Simple CSV: one column, one row per line of text.
    This is enough for Solace exports and easy to consume in Excel/Sheets.
    """
    buf = io.StringIO()
    writer = csv.writer(buf)

    for line in text.replace("\r", "").split("\n"):
        writer.writerow([line])

    return buf.getvalue().encode("utf-8")


def normalize_export_type(raw: str, body: Dict[str, Any]) -> str:
    """
    Normalize incoming export type, with a tiny bit of forgiveness:
    - "word" or "doc" -> "docx"
    - "excel" or "spreadsheet" -> "csv"
    - If type is empty but filename is provided, try to infer from extension.
    """
    raw = (raw or "").strip().lower()

    if not raw:
        filename = (body.get("filename") or "").lower()
        if filename.endswith(".pdf"):
            raw = "pdf"
        elif filename.endswith(".docx") or filename.endswith(".doc"):
            raw = "docx"
        elif filename.endswith(".csv"):
            raw = "csv"

    aliases = {
        "word": "docx",
        "doc": "docx",
        "docx": "docx",
        "pdf": "pdf",
        "csv": "csv",
        "excel": "csv",
        "spreadsheet": "csv",
    }

    return aliases.get(raw, raw)


# ------------------------------------------------------
# Basic routes (A + B: / and /healthz)
# ------------------------------------------------------
@app.get("/")
async def root() -> Dict[str, Any]:
    """
    Simple root endpoint so health checks and screenshots
    don't produce 404 spam in the logs.
    """
    return {
        "ok": True,
        "service": "mca-export-worker",
        "message": "Export worker is running.",
    }


@app.get("/healthz")
async def healthz() -> Dict[str, Any]:
    """
    Lightweight health probe endpoint.
    """
    return {
        "ok": True,
        "status": "healthy",
        "has_key": bool(PY_WORKER_KEY),
    }


# ------------------------------------------------------
# Main export endpoint
# ------------------------------------------------------
@app.post("/api/generate")
async def generate(request: Request):
    """
    Single entrypoint Solace calls from:
      PY_WORKER_URL = https://mca-export-worker.vercel.app/api/generate

    Expects JSON payload like:
      { "type": "pdf" | "docx" | "csv", "title": "...", "content": "..." }

    Returns raw file bytes with appropriate Content-Type.
    """

    # ---- Guard: ensure key is configured ----
    if not PY_WORKER_KEY:
        logger.error("PY_WORKER_KEY is not set in worker environment.")
        raise HTTPException(
            status_code=500,
            detail="Worker misconfigured: PY_WORKER_KEY is not set.",
        )

    # ---- Auth check ----
    auth_header = request.headers.get("authorization", "")
    expected = f"Bearer {PY_WORKER_KEY}"

    if auth_header != expected:
        logger.warning("Unauthorized request to /api/generate")
        raise HTTPException(status_code=401, detail="Unauthorized")

    # ---- Parse body safely (no Pydantic to avoid brittle 422s) ----
    try:
        body = await request.json()
    except Exception:
        logger.exception("Failed to parse JSON body for /api/generate")
        raise HTTPException(status_code=400, detail="Invalid JSON body")

    if not isinstance(body, dict):
        logger.error("Request body is not a JSON object: %r", body)
        raise HTTPException(
            status_code=400,
            detail="Request body must be a JSON object.",
        )

    raw_type = str(
        body.get("type")
        or body.get("format")
        or body.get("export_type")
        or ""
    )
    export_type = normalize_export_type(raw_type, body)

    title = (body.get("title") or "Solace Export").strip()
    content = body.get("content") or ""

    logger.info(
        "Export request: type=%s (raw=%s), title_len=%d, content_chars=%d",
        export_type,
        raw_type,
        len(title),
        len(content),
    )

    if not export_type:
        raise HTTPException(
            status_code=400,
            detail="Missing 'type' field (expected 'pdf', 'docx', or 'csv').",
        )
    if export_type not in {"pdf", "docx", "csv"}:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Unsupported type: '{export_type}'. "
                "Supported types are: 'pdf', 'docx', 'csv'."
            ),
        )

    try:
        # ---- Generate bytes based on type ----
        if export_type == "pdf":
            data = generate_pdf(title, content)
            media_type = "application/pdf"
            filename = "solace-export.pdf"
        elif export_type == "docx":
            data = generate_docx(title, content)
            media_type = (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            )
            filename = "solace-export.docx"
        else:  # "csv"
            data = generate_csv(content)
            media_type = "text/csv"
            filename = "solace-export.csv"

    except Exception:
        # E: Hardened error handling
        logger.exception("Error while generating %s export", export_type)
        raise HTTPException(
            status_code=500,
            detail="Export worker encountered an error while generating the file.",
        )

    headers = {
        # Let Solace / browser treat it as a download if this URL is hit directly
        "Content-Disposition": f'inline; filename="{filename}"',
        "X-MCA-Export-Type": export_type,
    }

    logger.info("Export generated successfully: type=%s, bytes=%d", export_type, len(data))

    return Response(content=data, media_type=media_type, headers=headers)
